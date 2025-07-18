# app/services/cv_writer.py

import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.opc.constants import RELATIONSHIP_TYPE as RT

class CVWriter:
    """
    Kage: The scribe of forms. This module is tasked with forging the
    Curriculum Vitae document, translating raw data into a structured,
    presentable form. It operates with precision, ensuring every detail
    finds its rightful place.
    """
    def __init__(self, output_dir="output"):
        """
        Kage: Initializes the CVWriter. A defined destination is essential
        for the manifestation of documents.
        """
        self.output_dir = output_dir
        os.makedirs(self.output_dir, exist_ok=True)
        print(f"[Kage CV] CVWriter initialized. Output directory: {self.output_dir}")

    def _add_heading(self, document, text, level=1, center=False):
        """Kage: Adds a section heading. Clarity in structure."""
        heading = document.add_heading(text, level=level)
        if center:
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _add_paragraph(self, document, text, bold=False, italic=False, font_size=None):
        """Kage: Adds a textual passage. Precision in presentation."""
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(2)
        run = paragraph.add_run(text)
        run.bold = bold
        run.italic = italic
        if font_size:
            run.font.size = Pt(font_size)
        return paragraph

    def _add_bullet_list(self, document, items):
        """Kage: Organizes information into a list. For clear enumeration."""
        for item in items:
            para = document.add_paragraph(item, style='List Bullet')
            para.paragraph_format.space_after = Pt(2)

    def _add_clickable_hyperlink(self, paragraph, url, text, underline=True):
        """Kage: Embeds a navigable path within the document. For external connections."""
        try:
            part = paragraph.part
            r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)

            hyperlink = OxmlElement('w:hyperlink')
            hyperlink.set(qn('r:id'), r_id)

            new_run = OxmlElement('w:r')
            r_pr = OxmlElement('w:rPr')

            if underline:
                u = OxmlElement('w:u')
                u.set(qn('w:val'), 'single')
                r_pr.append(u)

            color = OxmlElement('w:color')
            color.set(qn('w:val'), '0000FF')
            r_pr.append(color)

            text_elem = OxmlElement('w:t')
            text_elem.text = text
            new_run.append(text_elem)

            hyperlink.append(new_run)
            paragraph._element.append(hyperlink)
        except Exception as e:
            print(f"[Kage CV] Failed to add hyperlink '{text}': {e}")

    def _filter_technologies_for_display(self, technologies: list) -> list:
        """
        Kage: Filters out overly granular or less relevant technologies for CV display.
        Focuses on major frameworks, tools, and platforms, as requested by the Master.
        """
        # Define a list of common, often too granular, Python libraries or generic terms
        # that might be identified by LLM but are not ideal for high-level CV tech lists.
        # This list can be expanded based on feedback.
        EXCLUDE_TECH_KEYWORDS = [
            # Generic/Low-level Python libraries/tools
            "requests", "json", "os", "uuid", "random", "traceback", "dotenv",
            "smtplib", "email", "header", "urllib.parse", "httpx", "jwt",
            "uvicorn", "pygithub", "pyjwt", "pyyaml", "toml", "pypdf",
            "python-docx", "docx2pdf", "keyboard", "pygetwindow", "rasterio",
            "pyinstaller", "pyopengl", "pyqt5", # PyQt5 is a framework, but often listed with its components
            
            # Generic development concepts/features (better covered by skills/achievements)
            "api design", "authentication", "backend development", "frontend development",
            "github integration", "cv parsing", "data extraction", "llm-powered project analysis",
            "oauth", "project summarization", "resume generation", "scoring systems",
            "template design", "web ui development", "rest apis", "websocket",
            
            # Languages/Markup already covered in Programming Languages section (or to be combined)
            "css", "css3", "html", "html5", "javascript", "python", "bash", "c", "c++",
            "dart", "go", "java", "kotlin", "rust", "swift", "typescript", "tex",
            
            # Specific SDKs/components if the platform itself is listed
            "firebase js sdk", "firebase-admin", 
            
            # Other potentially too granular or implied tools
            "jinja2", # Template engine, can be implied by Python web frameworks
            "git", # Source control is a skill, not a specific tech for this list
            "linux", # Operating system, often implied by development environment
            "vs code", # IDE, generally not listed as a core project technology
            "excel", "google sheets", "power bi", "tableau", # Data analysis tools, but not core *project* tech
        ]
        
        # Normalize to lower case for comparison
        excluded_lower = [k.lower() for k in EXCLUDE_TECH_KEYWORDS]
        
        filtered = []
        for tech in technologies:
            if tech.lower() not in excluded_lower:
                filtered.append(tech)
        return filtered

    def generate_cv(self, projects_data: list, user_cv_data: dict, output_filename="resume.docx"):
        """
        Kage: Forges the complete CV document.
        It accepts all user-related data as a single dictionary (user_cv_data)
        and project data, then constructs the document with precision.

        Args:
            projects_data (list): A list of dictionaries, each representing an analyzed project.
            user_cv_data (dict): A dictionary containing all user profile information
                                 (name, email, phone, summary, skills, technologies, etc.).
            output_filename (str): The desired filename for the generated DOCX.

        Returns:
            str: The full path to the generated DOCX file, or None if generation fails.
        """
        document = Document()

        # Kage: Setting the document's voice. Compact and clear.
        style = document.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(10.5)

        # Kage: Destructuring the user's essence from the provided data.
        name = user_cv_data.get('name', 'Your Name')
        email = user_cv_data.get('email', '')
        phone = user_cv_data.get('phone', '')
        professional_summary = user_cv_data.get('professional_summary', '')
        user_defined_skills = user_cv_data.get('user_defined_skills', [])
        user_defined_technologies = user_cv_data.get('user_defined_technologies', [])
        languages = user_cv_data.get('languages', []) # This is for spoken languages
        linkedin = user_cv_data.get('linkedin', '')
        github_profile = user_cv_data.get('github_profile', '')
        work_experience = user_cv_data.get('work_experience', [])

        # Kage: Constructing the header. The identity.
        self._add_heading(document, name, level=1, center=True)
        contact_paragraph = document.add_paragraph()
        contact_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if linkedin and linkedin != "N/A":
            self._add_clickable_hyperlink(contact_paragraph, linkedin, "LinkedIn")
            contact_paragraph.add_run(" | ")
        if github_profile and github_profile != "N/A":
            self._add_clickable_hyperlink(contact_paragraph, github_profile, "GitHub")
            contact_paragraph.add_run(" | ")
        contact_items = []
        if email:
            contact_items.append(email)
        if phone and phone != "N/A":
            contact_items.append(phone)
        if contact_items:
            contact_paragraph.add_run(" | ".join(contact_items))

        # Kage: Adding the professional summary. The distilled purpose.
        if professional_summary:
            self._add_heading(document, "Professional Summary", level=2)
            self._add_paragraph(document, professional_summary)

        # Kage: Enumerating skills. The abilities honed.
        if user_defined_skills:
            self._add_heading(document, "Skills", level=2)
            self._add_paragraph(document, ", ".join(sorted(user_defined_skills)), bold=True)

        # Kage: Enumerating technologies. The tools mastered.
        if user_defined_technologies:
            self._add_heading(document, "Technologies", level=2)
            self._add_paragraph(document, ", ".join(sorted(user_defined_technologies)), bold=True)

        # Kage: Listing spoken languages. The communication paths.
        if languages:
            self._add_heading(document, "Languages", level=2)
            self._add_paragraph(document, ", ".join(sorted(languages)), bold=True)
            # Kage: This section is specifically for spoken languages (e.g., English, Arabic, Japanese).

        # Kage: Detailing work experience. The forged path.
        # Kage: Limiting work experience entries to a maximum of 2 for conciseness.
        max_work_experience_entries = 2 
        work_experience_to_display = work_experience[:max_work_experience_entries]

        if work_experience_to_display:
            self._add_heading(document, "Work Experience", level=2)
            for entry in work_experience_to_display:
                # Kage: Now correctly formatting the work experience entry from an object
                # into a readable paragraph.
                title = entry.get('title', 'N/A')
                company = entry.get('company', 'N/A')
                start_date = entry.get('start_date', 'N/A')
                end_date = entry.get('end_date', 'N/A')
                responsibilities = entry.get('responsibilities', [])

                # Add title and company
                work_exp_para = document.add_paragraph()
                work_exp_run = work_exp_para.add_run(f"{title} | {company}")
                work_exp_run.bold = True
                work_exp_run.font.size = Pt(11)

                # Add duration
                duration_text = f"{start_date} - {end_date}"
                self._add_paragraph(document, duration_text, italic=True)

                # Add responsibilities as bullet points
                if responsibilities:
                    self._add_bullet_list(document, responsibilities)
                
                # Add a small separator after each entry for readability
                document.add_paragraph().add_run("---").alignment = WD_ALIGN_PARAGRAPH.CENTER


        # Kage: Presenting projects. The demonstrated strength.
        if projects_data:
            self._add_heading(document, "Projects", level=2)
            # Sort projects by score in descending order for impact
            sorted_projects = sorted(projects_data, key=lambda p: p.get('score', 0), reverse=True)
            # The top 4 limit is already applied in main.py, so no need to slice here again.
            # projects_data already contains only the top 4.

            for project in sorted_projects:
                title_paragraph = document.add_paragraph()
                title_paragraph.paragraph_format.space_after = Pt(1)

                run = title_paragraph.add_run(project.get('name', 'Unnamed Project'))
                run.bold = True
                run.font.size = Pt(11)

                if project.get('html_url'):
                    title_paragraph.add_run(" | ")
                    self._add_clickable_hyperlink(title_paragraph, project['html_url'], "View Project")

                if project.get('summary'):
                    self._add_paragraph(document, project['summary'])

                if project.get('achievements'):
                    self._add_paragraph(document, "Key Achievements:", bold=True)
                    # Kage: Limit achievements to a maximum of 3 bullet points for conciseness.
                    self._add_bullet_list(document, project['achievements'][:3]) 

                # Kage: Combine Programming Languages and Key Technologies into a single filtered line.
                combined_proj_tech_lang = []

                # Add programming languages from GitHub data
                project_languages = project.get('languages', {})
                if project_languages:
                    # Filter out common web languages if they are the only ones, as per previous request
                    # This logic is now handled more robustly by the filter function if needed.
                    combined_proj_tech_lang.extend(project_languages.keys())

                # Add technologies identified by analyzer
                project_technologies = project.get('technologies', [])
                combined_proj_tech_lang.extend(project_technologies)
                
                # Apply filtering and then sort for final display
                filtered_combined_proj_tech_lang = self._filter_technologies_for_display(list(set(combined_proj_tech_lang)))
                
                if filtered_combined_proj_tech_lang: # Corrected variable name
                    self._add_paragraph(document, "Technologies & Languages: " + ", ".join(sorted(filtered_combined_proj_tech_lang)), italic=True)
                
        # Kage: Finalizing the form. Securing the manifestation.
        output_path = os.path.join(self.output_dir, output_filename)
        try:
            document.save(output_path)
            print(f"[Kage CV] CV generated successfully: {output_path}")
            return output_path
        except Exception as e:
            print(f"[Kage CV] Failed to save CV: {e}")
            return None

# Kage: Dummy version for testing purposes.
if __name__ == "__main__":
    print("[Kage CV Test] Running dummy CVWriter test...")

    # Dummy user data
    dummy_user_cv_data = {
        "name": "Ahmed Boray",
        "email": "silragonryu@example.com",
        "phone": "+1234567890",
        "linkedin": "https://www.linkedin.com/in/ahmed-ehab-992920344",
        "github_profile": "https://github.com/silragon-ryu",
        "professional_summary": "Disciplined and purpose-driven computer engineering student at Istinye University, deeply passionate about software engineering, AI, and embedded systems. Known for blending emotional intelligence with technical precision. Currently contributing to AI healthcare research and leading autonomous system design in national competitions.",
        "user_defined_skills": ["Authentic Communicator", "Calm Under Pressure", "Disciplined", "Emotionally Intelligent", "Grounded", "Leadership", "Mentor & Collaborator", "Purpose-Driven", "Strategic Thinker"],
        "user_defined_technologies": ["Django", "Docker", "FastAPI", "Firebase", "Flask", "Flutter", "Hugging Face", "Jupyter", "MongoDB", "MySQL", "Node.js", "NumPy", "OpenCV", "Pandas", "PostgreSQL", "PyTorch", "React", "Scikit-learn", "TensorFlow", "Tailwind", "WebSocket"],
        "languages": ["English (Native)", "Arabic (Fluent)", "Japanese (Basic)"], # Human spoken languages
        "work_experience": [
            {
                "title": "Senior Software Engineer",
                "company": "Tech Solutions Inc.",
                "start_date": "Jan 2022",
                "end_date": "Present",
                "responsibilities": [
                    "Led development of microservices using Python and Django.",
                    "Optimized database queries, reducing response times by 30%.",
                    "Mentored junior developers and conducted code reviews."
                ]
            },
            {
                "title": "Software Developer",
                "company": "Innovate Corp.",
                "start_date": "Feb 2019",
                "end_date": "Dec 2021",
                "responsibilities": [
                    "Developed and maintained critical backend systems with Node.js.",
                    "Collaborated on front-end features using React.",
                    "Implemented new API endpoints for mobile integration."
                ]
            },
            {
                "title": "Junior Developer Intern",
                "company": "Startup X",
                "start_date": "Sum 2018",
                "end_date": "Sum 2018",
                "responsibilities": [
                    "Assisted in front-end development using HTML and CSS.",
                    "Participated in daily stand-ups and code reviews."
                ]
            }
        ]
    }

    # Dummy project data (top 4, sorted by a dummy score)
    dummy_projects_data = [
        {
            "name": "kaku-ryu",
            "html_url": "https://github.com/silragon-ryu/kaku-ryu",
            "summary": "Kaku-Ryu is a dynamic resume engine that automatically updates a professional CV by analyzing GitHub repositories using a local LLM (Llama 3). It generates polished DOCX and PDF resumes with AI-powered project insights and scoring.",
            "achievements": [
                "Dynamic resume generation from GitHub repositories",
                "AI-powered project insights and scoring",
                "Generation of DOCX and PDF resumes",
                "Integration with local LLM (Llama 3) for code analysis"
            ],
            "languages": {"Python": 80, "JavaScript": 15, "HTML": 5},
            "technologies": ["FastAPI", "Llama 3", "Ollama", "Firebase", "Tailwind CSS"],
            "score": 95.5
        },
        {
            "name": "RYU-Scape",
            "html_url": "https://github.com/silragon-ryu/RYU-Scape",
            "summary": "RYU-Scape is a desktop application that transforms 2D satellite imagery into detailed 3D scenes using deep learning. It combines RGB images with DSM data to reconstruct interactive, textured environments with procedural details.",
            "achievements": [
                "Developed an AI-powered 2D-to-3D converter for satellite images.",
                "Implemented DeepLabV3 model for building footprint detection.",
                "Automated DSM handling for seamless data integration."
            ],
            "languages": {"Python": 90, "Jupyter Notebook": 10},
            "technologies": ["PyTorch", "OpenCV", "Open3D", "DeepLabV3"],
            "score": 92.0
        },
        {
            "name": "focus_lock",
            "html_url": "https://github.com/silragon-ryu/focus_lock",
            "summary": "focus_lock is a Python-based desktop application designed to enhance focus and productivity by eliminating distractions. It integrates with SumatraPDF and Spotify, blocks system shortcuts, hides the taskbar, and provides a session timer and lockdown feature.",
            "achievements": [
                "Developed a desktop application for enhanced focus and productivity.",
                "Integrated with SumatraPDF for distraction-free PDF reading.",
                "Implemented aggressive shortcut blocking to minimize interruptions."
            ],
            "languages": {"Python": 95, "HTML": 5},
            "technologies": ["PyQt5", "Spotify API"],
            "score": 88.0
        },
        {
            "name": "felixx",
            "html_url": "https://github.com/silragon-ryu/felixx",
            "summary": "FelixxCars is a modern and responsive car rental platform built with HTML5, CSS3, and JavaScript using Bootstrap 5. It offers a seamless booking experience with features like dark/light mode toggling, interactive car listings, and a customer-focused design.",
            "achievements": [
                "Developed a responsive car rental platform.",
                "Implemented dark/light mode toggling with local storage.",
                "Created interactive car listings with detailed specifications."
            ],
            "languages": {"HTML": 40, "CSS": 30, "JavaScript": 30},
            "technologies": ["Bootstrap 5", "Tailwind CSS"],
            "score": 85.0
        },
        { # This project should not be included as it's outside the top 4
            "name": "low-score-project",
            "html_url": "https://github.com/silragon-ryu/low-score-project",
            "summary": "A simple utility project.",
            "achievements": ["Implemented basic functionality."],
            "languages": {"Python": 100},
            "technologies": ["requests"],
            "score": 50.0
        }
    ]

    # Sort dummy projects to simulate main.py's behavior of selecting top 4
    dummy_projects_data_sorted_top4 = sorted(dummy_projects_data, key=lambda p: p.get('score', 0), reverse=True)[:4]

    cv_writer = CVWriter()
    output_filename = "dummy_test_resume.docx"
    generated_path = cv_writer.generate_cv(dummy_projects_data_sorted_top4, dummy_user_cv_data, output_filename)

    if generated_path:
        print(f"[Kage CV Test] Dummy CV generated at: {generated_path}")
        print("Please check the 'output' directory for 'dummy_test_resume.docx'.")
    else:
        print("[Kage CV Test] Failed to generate dummy CV.")
