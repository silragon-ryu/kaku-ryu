# app/services/cv_parser.py

import os
import json
import requests
from dotenv import load_dotenv
import re
from docx import Document
from pypdf import PdfReader # For PDF text extraction

# Load environment variables
load_dotenv()

class CVParser:
    def __init__(self, api_url="https://generativelanguage.googleapis.com/v1beta/models/gemini-2.0-flash:generateContent", model_name="gemini-2.0-flash", api_key: str = None):
        """
        Initializes the CVParser with Google Gemini API configuration.
        Args:
            api_url (str): The base URL for the Gemini API's generateContent endpoint.
            model_name (str): The name of the Gemini model to use.
            api_key (str, optional): The Gemini API key. If None, it falls back to GEMINI_API environment variable.
        """
        self.api_url = api_url
        self.model_name = model_name
        self.api_key = api_key if api_key else os.getenv("GEMINI_API", "")

        print(f"[Kage CV Parser] CVParser init: GEMINI_API key loaded: {'(not set or empty)' if not self.api_key else '*****' + self.api_key[-4:]}")

        if not self.api_key:
            raise ValueError(
                "[Kage CV Parser] GEMINI_API key is not set. Please ensure you have a valid "
                "'GEMINI_API' entry in your .env file in the root directory, "
                "or that the environment variable is otherwise provided."
            )
        print(f"[Kage CV Parser] CVParser initialized for Gemini model: {self.model_name}")

    def _extract_text_from_docx(self, docx_path: str) -> str:
        """Extracts text from a .docx file."""
        try:
            document = Document(docx_path)
            full_text = []
            for para in document.paragraphs:
                full_text.append(para.text)
            return "\n".join(full_text)
        except Exception as e:
            print(f"Error extracting text from DOCX {docx_path}: {e}")
            return ""

    def _extract_text_from_pdf(self, pdf_path: str) -> str:
        """Extracts text from a .pdf file using pypdf."""
        try:
            reader = PdfReader(pdf_path)
            full_text = []
            for page in reader.pages:
                full_text.append(page.extract_text())
            return "\n".join(full_text)
        except Exception as e:
            print(f"Error extracting text from PDF {pdf_path}: {e}")
            return ""

    def _sanitize_text(self, text: str) -> str:
        """
        Sanitizes text to remove or replace problematic Unicode characters that might cause
        "Bad Request" errors with LLM APIs.
        Encodes to ASCII, ignoring unencodable characters, then decodes back.
        """
        if not isinstance(text, str):
            return ""

        sanitized_text = text.encode('ascii', 'ignore').decode('ascii')
        sanitized_text = sanitized_text.replace('“', '"').replace('”', '"')
        sanitized_text = sanitized_text.replace('‘', "'").replace('’', "'")
        sanitized_text = sanitized_text.replace('—', '--').replace('–', '-')
        sanitized_text = sanitized_text.replace('…', '...')
        sanitized_text = sanitized_text.replace('\u2013', '-') # En dash
        sanitized_text = sanitized_text.replace('\u2014', '--') # Em dash
        sanitized_text = sanitized_text.replace('\u2019', "'") # Right single quotation mark
        sanitized_text = sanitized_text.replace('\u201c', '"') # Left double quotation mark
        sanitized_text = sanitized_text.replace('\u201d', '"') # Right double quotation mark
        sanitized_text = sanitized_text.replace('\u2022', '*') # Bullet point
        sanitized_text = sanitized_text.replace('\u00a0', ' ') # Non-breaking space
        sanitized_text = re.sub(r'[\x00-\x1F\x7F]+', '', sanitized_text) # Remove non-printable ASCII
        return sanitized_text.strip()

    def parse_cv(self, cv_path: str) -> dict:
        """
        Parses a CV file (DOCX or PDF) to extract structured information using Gemini.
        Args:
            cv_path (str): The path to the CV file.
        Returns:
            dict: A dictionary containing parsed CV data (name, email, phone, linkedin,
                  professional_summary, user_defined_skills, user_defined_technologies,
                  languages, work_experience).
        """
        file_extension = os.path.splitext(cv_path)[1].lower()
        cv_text = ""

        if file_extension == ".docx":
            cv_text = self._extract_text_from_docx(cv_path)
        elif file_extension == ".pdf":
            cv_text = self._extract_text_from_pdf(cv_path)
        else:
            raise ValueError("Unsupported file type. Only .docx and .pdf are supported.")

        if not cv_text:
            print(f"[Kage CV Parser] No text extracted from {cv_path}.")
            return self._fallback_data()

        # Sanitize the extracted text before sending to the LLM
        sanitized_cv_text = self._sanitize_text(cv_text)
        
        # Limit the CV text length to prevent excessively long prompts
        MAX_CV_TEXT_LENGTH = 20000 # Characters
        if len(sanitized_cv_text) > MAX_CV_TEXT_LENGTH:
            sanitized_cv_text = sanitized_cv_text[:MAX_CV_TEXT_LENGTH] + "\n... (truncated CV content)"
            print(f"[Kage CV Parser] CV text truncated for LLM processing.")

        prompt = f"""
You are an expert CV parser. Analyze the following CV text and extract the information into a structured JSON format. Provide no explanations, no markdown, only the JSON. Ensure the JSON is clean, valid, and immediately usable.

CV Text:
---
{sanitized_cv_text}
---

Extract the following fields:
- "name": Full name of the candidate.
- "email": Candidate's email address.
- "phone": Candidate's phone number.
- "linkedin": URL of the candidate's LinkedIn profile.
- "professional_summary": A concise professional summary (1-3 sentences) from the CV.
- "user_defined_skills": A list of key skills mentioned in the CV (e.g., "Project Management", "Data Analysis", "Cloud Computing").
- "user_defined_technologies": A list of specific technologies/tools mentioned (e.g., "Python", "React", "AWS", "Docker", "TensorFlow").
- "spoken_languages": A list of human spoken languages mentioned (e.g., "English (Fluent)", "Spanish (Conversational)").
- "work_experience": An array of objects, each representing a work experience entry. Each object should have:
    - "title": Job title.
    - "company": Company name.
    - "start_date": Start date of employment (e.g., "Jan 2020", "2019").
    - "end_date": End date of employment (e.g., "Dec 2022", "Present", "2021").
    - "responsibilities": A list of key responsibilities/achievements for that role.

If a field is not found or applicable, use "N/A" for strings, empty list [] for arrays, or empty object {{}} for objects.

Return JSON only:
{{
  "name": "N/A",
  "email": "N/A",
  "phone": "N/A",
  "linkedin": "N/A",
  "professional_summary": "N/A",
  "user_defined_skills": [],
  "user_defined_technologies": [],
  "spoken_languages": [],
  "work_experience": [
    {{
      "title": "N/A",
      "company": "N/A",
      "start_date": "N/A",
      "end_date": "N/A",
      "responsibilities": []
    }}
  ]
}}
"""
        payload = {
            "contents": [
                {
                    "role": "user",
                    "parts": [{"text": prompt}]
                }
            ],
            "generationConfig": {
                "responseMimeType": "application/json",
                "responseSchema": {
                    "type": "OBJECT",
                    "properties": {
                        "name": {"type": "STRING"},
                        "email": {"type": "STRING"},
                        "phone": {"type": "STRING"},
                        "linkedin": {"type": "STRING"},
                        "professional_summary": {"type": "STRING"},
                        "user_defined_skills": {"type": "ARRAY", "items": {"type": "STRING"}},
                        "user_defined_technologies": {"type": "ARRAY", "items": {"type": "STRING"}},
                        "spoken_languages": {"type": "ARRAY", "items": {"type": "STRING"}},
                        "work_experience": {
                            "type": "ARRAY",
                            "items": {
                                "type": "OBJECT",
                                "properties": {
                                    "title": {"type": "STRING"},
                                    "company": {"type": "STRING"},
                                    "start_date": {"type": "STRING"},
                                    "end_date": {"type": "STRING"},
                                    "responsibilities": {"type": "ARRAY", "items": {"type": "STRING"}}
                                },
                                "required": ["title", "company", "start_date", "end_date", "responsibilities"]
                            }
                        }
                    },
                    "required": ["name", "email", "phone", "linkedin", "professional_summary", "user_defined_skills", "user_defined_technologies", "spoken_languages", "work_experience"]
                },
                "temperature": 0.2,
                "topP": 0.9
            }
        }

        text = "" # Initialize text for error messages

        try:
            print(f"[Kage CV Parser] Sending CV text to Gemini for parsing...")
            api_url_with_key = f"{self.api_url}?key={self.api_key}"
            
            res = requests.post(api_url_with_key, headers={'Content-Type': 'application/json'}, json=payload, timeout=180)
            res.raise_for_status()

            gemini_response = res.json()
            if gemini_response.get("candidates") and len(gemini_response["candidates"]) > 0 and \
               gemini_response["candidates"][0].get("content") and \
               gemini_response["candidates"][0]["content"].get("parts") and \
               len(gemini_response["candidates"][0]["content"]["parts"]) > 0:
                text = gemini_response["candidates"][0]["content"]["parts"][0].get("text", "")
            else:
                print(f"[Kage CV Parser] ❌ Gemini response structure unexpected: {gemini_response}")
                return self._fallback_data()

            parsed_data = json.loads(text)
            print(f"[Kage CV Parser] CV parsed successfully.")
            
            final_parsed_data = {
                "name": parsed_data.get("name", "N/A"),
                "email": parsed_data.get("email", "N/A"),
                "phone": parsed_data.get("phone", "N/A"),
                "linkedin": parsed_data.get("linkedin", "N/A"),
                "github_profile": parsed_data.get("github_profile", "N/A"),
                "professional_summary": parsed_data.get("professional_summary", "N/A"),
                "user_defined_skills": parsed_data.get("user_defined_skills", []),
                "user_defined_technologies": parsed_data.get("user_defined_technologies", []),
                "languages": parsed_data.get("spoken_languages", []),
                "work_experience": parsed_data.get("work_experience", []) # Now expects list of objects
            }
            return final_parsed_data

        except requests.exceptions.ConnectionError:
            print(f"[Kage CV Parser] ❌ Cannot reach Gemini API. Check network connection.")
        except requests.exceptions.Timeout:
            print(f"[Kage CV Parser] ❌ Timeout during CV parsing with Gemini API.")
        except requests.exceptions.RequestException as e:
            error_details = f"Status Code: {e.response.status_code if e.response else 'N/A'}"
            if e.response and e.response.text:
                error_details += f", Response Body: {e.response.text}"
            print(f"[Kage CV Parser] ❌ HTTP error with Gemini API: {e}. Details: {error_details}")
            print(f"Gemini Raw Output (if available): {text[:500]}...")
        except json.JSONDecodeError:
            print(f"[Kage CV Parser] ❌ JSON parsing failed from Gemini API. Output:\n{text}")
        except Exception as e:
            print(f"[Kage CV Parser] ❌ Unexpected error during CV parsing: {e}")

        return self._fallback_data()

    def _fallback_data(self) -> dict:
        """Returns a default structure in case of parsing failure."""
        print("[Kage CV Parser] Using fallback data for CV parsing.")
        return {
            "name": "N/A",
            "email": "N/A",
            "phone": "N/A",
            "linkedin": "N/A",
            "github_profile": "N/A",
            "professional_summary": "Could not extract professional summary.",
            "user_defined_skills": [],
            "user_defined_technologies": [],
            "languages": [],
            "work_experience": [] # Fallback for list of objects
        }

# Example Usage:
if __name__ == "__main__":
    # Create a dummy DOCX file for testing
    dummy_docx_content = """
    John Doe
    john.doe@example.com | +1-555-123-4567 | LinkedIn: linkedin.com/in/johndoe

    Professional Summary
    Highly motivated software engineer with 5 years of experience in developing scalable web applications using Python and JavaScript. Passionate about creating efficient and user-friendly solutions.

    Skills
    Project Management, Data Analysis, Cloud Computing, Problem Solving, Communication

    Technologies
    React, AWS, Docker, TensorFlow

    Languages
    English (Fluent), Spanish (Conversational), French (Basic)

    Work Experience
    Senior Software Engineer | Tech Solutions Inc. | Jan 2022 - Present
    - Led development of microservices using Python and Django.
    - Optimized database queries, reducing response times by 30%.
    - Mentored junior developers and conducted code reviews.

    Software Developer | Innovate Corp. | Feb 2019 - Dec 2021
    - Developed and maintained critical backend systems with Node.js.
    - Collaborated on front-end features using React.
    - Implemented new API endpoints for mobile integration.

    Education
    B.S. in Computer Science | University of Example | 2019
    """
    
    # Create a dummy PDF file for testing (requires reportlab or similar for actual PDF creation)
    # For a real test, you'd need to convert the DOCX to PDF or have a pre-existing PDF.
    # Here, we'll just simulate a PDF path and content.
    dummy_pdf_content = """
    Jane Smith
    jane.smith@example.com | +1-111-222-3333 | LinkedIn: linkedin.com/in/janesmith

    Summary
    Data Scientist with 3 years of experience in machine learning, data analysis, and predictive modeling. Proficient in Python and R.

    Key Skills
    Machine Learning, Deep Learning, Data Visualization, Statistical Analysis, NLP

    Tools
    Python, R, Pandas, NumPy, Scikit-learn, Matplotlib, Seaborn, TensorFlow, Keras

    Languages
    English (Native), French (Intermediate)

    Experience
    Data Scientist | Data Insights Co. | Mar 2021 - Present
    - Built predictive models for customer segmentation.
    - Developed dashboards for key business metrics.
    - Performed A/B testing and analyzed results.

    Junior Data Analyst | Analytics Hub | Jul 2019 - Feb 2021
    - Cleaned and preprocessed large datasets.
    - Created reports and presentations for stakeholders.
    """

    dummy_docx_path = "output/dummy_cv.docx"
    dummy_pdf_path = "output/dummy_cv.pdf" # This would be a real PDF file

    # Save dummy DOCX
    try:
        doc = Document()
        doc.add_paragraph(dummy_docx_content)
        doc.save(dummy_docx_path)
        print(f"Dummy DOCX saved to {dummy_docx_path}")
    except Exception as e:
        print(f"Could not create dummy DOCX: {e}")
        dummy_docx_path = None

    parser = CVParser()

    if dummy_docx_path and os.path.exists(dummy_docx_path):
        print("\n--- Parsing Dummy DOCX CV ---")
        parsed_docx_data = parser.parse_cv(dummy_docx_path)
        print(json.dumps(parsed_docx_data, indent=2))
        os.remove(dummy_docx_path) # Clean up dummy file
